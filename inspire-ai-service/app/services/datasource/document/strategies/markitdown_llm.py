import io
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Optional
import pymupdf
from markitdown import MarkItDown
from openai import OpenAI
from llama_index.core.schema import Document

from app.settings import settings
from app.services.datasource.document.strategies.base import DocumentConversionStrategy
from app.services.media_storage import MediaStorageService



class MarkItDownLLMStrategy(DocumentConversionStrategy):
    """Strategy for converting files using MarkItDown with LLM support"""

    def __init__(self):
        # API base of the Google that support for the openai library
        API_BASE = "https://generativelanguage.googleapis.com/v1beta/openai/"

        llm_client = OpenAI(
            base_url=API_BASE,
            api_key=settings.llm.ocr_gemini.api_key,
        )
        llm_model = settings.llm.ocr_gemini.model_name

        self.md = MarkItDown(llm_client=llm_client, llm_model=llm_model, temperature=settings.llm.ocr_gemini.temperature)
        self.media_storage = MediaStorageService()

    def _process_pdf_page(self, doc, page_number):
        """Process a single PDF page and return its content with page number for ordering"""
        page = doc.load_page(page_number)

        # Extract links from the page
        links = page.get_links()
        link_text = ""
        if links:
            for link in links:
                if "uri" in link:
                    link_text += f"\nLink: {link['uri']}"

        # Get page content as image
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")
        img_stream = io.BytesIO(img_bytes)

        # Convert image to text using MarkItDown with LLM
        document_content = self.md.convert_stream(img_stream, file_extension="png")

        # Save the page image if media storage is provided
        image_url_text = ""
        if self.media_storage:
            try:
                image_url = self.media_storage.save_image(
                    image_bytes=img_bytes,
                    page_num=page_number + 1,
                    description=document_content.text_content[:500] if document_content.text_content else None
                )
                # Inject image URL directly into the text
                if image_url:
                    page_desc = document_content.text_content[:100] if document_content.text_content else f"Page {page_number + 1}"
                    image_url_text = f"\n{page_desc}: {image_url}\n"
            except Exception:
                pass

        # Combine text content with image URL and links
        full_content = document_content.text_content
        if image_url_text:
            full_content += "\n" + image_url_text
        if link_text:
            full_content += "\n" + link_text

        return page_number, full_content, None

    def convert_pdf(self, file_path: str, metadata: dict = {}) -> list[Document]:
        """Special handling for PDF files using parallel processing"""
        doc = pymupdf.open(file_path)
        page_contents = {}

        with ThreadPoolExecutor(max_workers=min(2, len(doc))) as executor:
            future_to_page = {
                executor.submit(self._process_pdf_page, doc, page_num): page_num
                for page_num in range(len(doc))
            }

            for future in as_completed(future_to_page):
                try:
                    page_number, content, _ = future.result()
                    page_contents[page_number] = content
                except Exception:
                    page_num = future_to_page[future]
                    page_contents[page_num] = ""

        # Combine pages in correct order with inline image URLs
        final_content = ""
        for page_number in range(len(doc)):
            final_content += page_contents.get(page_number, "")

        doc.close()
        return [Document(text=final_content, metadata=metadata)]

    def convert(self, file_path: str, metadata: dict = {}) -> list[Document]:
        file_extension = Path(file_path).suffix.lower()

        # Special handling for PDFs due to their complexity
        if file_extension == '.pdf':
            return self.convert_pdf(file_path, metadata)

        # Special handling for DOCX files with images
        if file_extension == '.docx' and self.media_storage:
            return self.convert_docx(file_path, metadata)

        # For other file types that need LLM processing
        result = self.md.convert(file_path)
        return [Document(text=result.text_content, metadata=metadata)]

    def _extract_text_context(self, doc, image_rId: str, context_paragraphs: int = 3) -> tuple[str, str, str]:
        """
        Extract text context around an image in the document.

        Args:
            doc: DocxDocument instance
            image_rId: Relationship ID of the image
            context_paragraphs: Number of paragraphs to extract before and after

        Returns:
            Tuple of (before_text, caption_text, after_text)
        """
        before_text = ""
        after_text = ""
        caption_text = ""

        # Find paragraphs/runs that contain this image
        image_found = False
        para_texts = []

        for para in doc.paragraphs:
            para_text = para.text.strip()

            # Check if this paragraph contains the image
            has_image = False
            for run in para.runs:
                if hasattr(run, '_element'):
                    # Check for drawing elements (images)
                    for drawing in run._element.findall('.//{*}drawing'):
                        for blip in drawing.findall('.//{*}blip'):
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed == image_rId:
                                has_image = True
                                image_found = True
                                # Try to get caption from the same paragraph or nearby
                                if para_text:
                                    caption_text = para_text
                                break

            if has_image:
                # Get text before image
                if len(para_texts) > 0:
                    before_text = "\n".join(para_texts[-context_paragraphs:])
                continue

            if image_found and len(after_text.split("\n")) < context_paragraphs:
                # Collect paragraphs after image
                if para_text:
                    if after_text:
                        after_text += "\n" + para_text
                    else:
                        after_text = para_text
            elif not image_found:
                # Collect paragraphs before image
                if para_text:
                    para_texts.append(para_text)

        return before_text, caption_text, after_text

    def _generate_image_description(self, before_text: str, caption_text: str, after_text: str, image_counter: int) -> str:
        """
        Generate a contextual description for an image using LLM.

        Args:
            before_text: Text appearing before the image
            caption_text: Caption or text in the same paragraph as image
            after_text: Text appearing after the image
            image_counter: Image number in document

        Returns:
            Generated description string
        """
        try:
            # Build context prompt
            context_parts = []
            if before_text:
                context_parts.append(f"Text before image:\n{before_text[:500]}")
            if caption_text:
                context_parts.append(f"Image caption/label:\n{caption_text[:200]}")
            if after_text:
                context_parts.append(f"Text after image:\n{after_text[:500]}")

            if not context_parts:
                return f"Image {image_counter} from document"

            context = "\n\n".join(context_parts)

            # Use LLM to generate description
            prompt = f"""Based on the surrounding text context, generate a concise description (2-3 sentences) of what this image likely shows or illustrates.

{context}

Description:"""

            response = self.md.llm_client.chat.completions.create(
                model=self.md.llm_model,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that generates concise image descriptions based on document context."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=150
            )

            description = response.choices[0].message.content.strip()

            # Add contextual information
            full_description = description
            if caption_text:
                full_description = f"{caption_text} | {description}"

            return full_description

        except Exception:
            # Fallback to simple description with available context
            if caption_text:
                return f"{caption_text} (Image {image_counter})"
            return f"Image {image_counter} from document"

    def convert_docx(self, file_path: str, metadata: dict = {}) -> list[Document]:
        """Special handling for DOCX files with context-aware image extraction"""


        try:
            from docx import Document as DocxDocument
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except ImportError:
            result = self.md.convert(file_path)
            return [Document(text=result.text_content, metadata=metadata)]

        try:
            doc = DocxDocument(file_path)
            image_markers = {}  # Map rId to image info for inline markers

            # Extract all images from the document with context
            image_counter = 0

            # 1) Relationship-based detection with context extraction
            for rel_id, rel in doc.part.rels.items():
                try:
                    reltype = str(rel.reltype)
                    target_ref = getattr(rel, 'target_ref', '') or ''
                    if (hasattr(RT, 'IMAGE') and reltype == RT.IMAGE) or ('image' in target_ref):
                        image_counter += 1
                        image_bytes = rel.target_part.blob

                        # Extract context around this image
                        before_text, caption_text, after_text = self._extract_text_context(doc, rel_id)

                        # Generate contextual description
                        description = self._generate_image_description(
                            before_text, caption_text, after_text, image_counter
                        )

                        if self.media_storage:
                            try:
                                image_url = self.media_storage.save_image(
                                    image_bytes=image_bytes,
                                    page_num=image_counter,
                                    description=description
                                )
                                image_markers[rel_id] = {
                                    "description": description,
                                    "url": image_url,
                                    "counter": image_counter
                                }
                            except Exception:
                                pass
                except Exception:
                    pass

            # 2) Build text with inline image URLs at their natural positions
            # Walk through document and insert image URLs where images appear
            text_parts = []
            for para in doc.paragraphs:
                para_text = para.text.strip()

                # Add paragraph text first if it exists
                if para_text:
                    text_parts.append(para_text)

                # Check if paragraph contains images and inject URL inline
                for run in para.runs:
                    if hasattr(run, '_element'):
                        for drawing in run._element.findall('.//{*}drawing'):
                            for blip in drawing.findall('.//{*}blip'):
                                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed in image_markers:
                                    img_info = image_markers[embed]
                                    # Inject the actual URL inline in the text
                                    image_citation = f"{img_info['description']}: {img_info['url']}"
                                    text_parts.append(image_citation)
                                    break

            # Use the enhanced text with inline image citations
            text_content = "\n".join(text_parts) if text_parts else ""

            # Return document with image URLs embedded directly in text
            return [Document(text=text_content, metadata=metadata)]

        except Exception:
            # Fallback to basic conversion
            result = self.md.convert(file_path)
            return [Document(text=result.text_content, metadata=metadata)]
